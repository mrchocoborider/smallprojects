import re
import docx
from docx.shared import Pt
import copy


reg = re.compile(r'[a-zA-Z]+')

#Most of these methods will have to be customized to new documents
#Be aware of that fact!

def reOrder(docname):
    #docname should be a string
    doc = docx.Document(docname)

    fullText = []
    ctxt = []
    etxt = []

    #make a variable called lastline to see if the last line
    #was English or Chinese (they need to be kept together)
    lastLine = ""

    #we're going to switch the order of English and Chinese lines
    k = 0
    for i in doc.paragraphs:
        #First two lines should be appended without change
        if k == 0 or k == 1:
            fullText.append(i)
        if k != 0 and k != 1:
            #English
            if reg.search(i.text) != None:
                thisLine = "eng"
                #I was overthinking it before, should be good now, with the exception of the first 10 lines or so
                #which I just reordered manually, I just needed to put the Chinese on hold and then paste the English when
                #I got to it, and paste the Chinese that was stored up after that.
                if lastLine == "chi" and thisLine == "eng":
                    #if len(etxt) == 0:
                    #    fullText.append(i)

                    #append English first
                    fullText.append(i)
                    #if len(etxt) > 0:
                    #    for e in etxt:
                    #        fullText.append(e)
                    #        if k <= 10:
                    #            print e.text
                    #append Chinese next
                    if len(ctxt) > 0:
                        for c in ctxt:
                            fullText.append(c)
                            if k <= 10:
                                print c.text
                    #clear the lists
                    etxt = []
                    ctxt = []

                    #etxt.append(i)
                    lastLine = "eng"
                #Still not sure if there will be more english!
                elif lastLine == "eng" and thisLine == "eng":
                    etxt.append(i)
                    lastLine = "eng"
                
            #Chinese
            if reg.search(i.text) == None:
                thisLine = "chi"
                #append first line of Chinese and append to ctxt
                if lastLine == "" and thisLine == "chi":
                    ctxt.append(i)
                    lastLine = "chi"
                #if this line is Chinese and last line was Eng
                #there may be Chinese and Eng in the queue that we need
                #to append to fulltext before continuing 
                elif thisLine == "chi" and lastLine == "eng":
                    #append English first
                    if len(etxt) > 0:
                        for e in etxt:
                            fullText.append(e)
                    #append Chinese next
                    if len(ctxt) > 0:
                        for c in ctxt:
                            fullText.append(c)
                    #clear the lists
                    etxt = []
                    ctxt = []
                    #put this line in Chinese list
                    ctxt.append(i)

                    lastLine = "chi"
                #append the chinese to ctxt and skip until a language change
                elif thisLine == "chi" and lastLine == "chi":
                    ctxt.append(i)
                    lastLine = "chi"



        k += 1

    #blah = '\n'.join(fullText)
    #print blah

    new = docx.Document()

    for j in fullText:
        new.add_paragraph(j.text)
    #an attempt at copying the paragraph objects directly
    #for j in fullText:
    #    new.paragraphs.append(copy.deepcopy(j))

    #print new.paragraphs

    new.save('testing.docx')

                            
                            
doc.save('vanessa2.docx')



reOrder('reorder.docx')
#fontSize('testing.docx')
#delEmpty('vanessa.docx')
