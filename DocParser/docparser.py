import re
import docx
from docx.shared import Pt
import copy
import sys


reg = re.compile(r'[a-zA-Z]+')
#This was written specifically to help my fiancee who had huge documents of text in
#Chinese with English translation, but wanted to remove empty lines, change English font size to 11
#and also change the order of English and Chinese lines.
#Most of these methods will have to be customized to new documents
#Be aware of that fact!

def reOrder(doc1, doc2):
    #docname should be a string
    doc = docx.Document(doc1)

    fullText = []
    ctxt = []
    etxt = []

    #make a variable called lastline to see if the last line
    #was English or Chinese (they need to be kept together)
    lastLine = ""

    #we're going to switch the order of English and Chinese lines
    #k = 0
    for k, i in enumerate(doc.paragraphs):
        #First two lines should be appended without change
        if k == 0 or k == 1:
            fullText.append(i)
        if k != 0 and k != 1:
            #English
            if reg.search(i.text) != None:
                thisLine = "eng"
                #I was overthinking it before, should be good now, with the exception of the first 10 lines or so
                #which I just reordered manually, I just needed to put the Chinese on hold and then paste the English when
                #I got to it, and paste the Chinese that was stored up after that.
                if lastLine == "chi" and thisLine == "eng":
                    #append English first
                    fullText.append(i)
                    #append Chinese next
                    if len(ctxt) > 0:
                        for c in ctxt:
                            fullText.append(c)
                            if k <= 10:
                                #print(c.text)
                    #clear the lists
                    etxt = []
                    ctxt = []

                    #etxt.append(i)
                    lastLine = "eng"
                #Still not sure if there will be more english!
                elif lastLine == "eng" and thisLine == "eng":
                    etxt.append(i)
                    lastLine = "eng"
                
            #Chinese
            if reg.search(i.text) == None:
                thisLine = "chi"
                #append first line of Chinese and append to ctxt
                if lastLine == "" and thisLine == "chi":
                    ctxt.append(i)
                    lastLine = "chi"
                #if this line is Chinese and last line was Eng
                #there may be Chinese and Eng in the queue that we need
                #to append to fulltext before continuing 
                elif thisLine == "chi" and lastLine == "eng":
                    #append English first
                    if len(etxt) > 0:
                        for e in etxt:
                            fullText.append(e)
                    #append Chinese next
                    if len(ctxt) > 0:
                        for c in ctxt:
                            fullText.append(c)
                    #clear the lists
                    etxt = []
                    ctxt = []
                    #put this line in Chinese list
                    ctxt.append(i)

                    lastLine = "chi"
                #append the chinese to ctxt and skip until a language change
                elif thisLine == "chi" and lastLine == "chi":
                    ctxt.append(i)
                    lastLine = "chi"


        #k += 1


    new = docx.Document()

    for j in fullText:
        new.add_paragraph(j.text)

    new.save(doc2)



#this is for changing font size of English only lines
def fontSize(doc1, doc2):
    #docname should be a string
    doc = docx.Document(doc1)
       
    for i in doc.paragraphs:
        #we have to access all the run objects to get to font size
        #for the first doc, each line was a new paragraph
        #another doc needs to break it down to runs
        for j in i.runs:
            if reg.search(j.text) != None:
                j.font.size = Pt(11)
                                
                            
    doc.save(doc2)


#delete all blank lines in the document
def delEmpty(doc1, doc2):
    #This is written for a specific file that has some weird 
    #characteristics. All the text is organized into just eleven
    #paragraphs
    doc = docx.Document(doc1)
    new = docx.Document()

    fullText = []

    #print len(doc.paragraphs)
    for i in doc.paragraphs:
        fullText.append(i.text)

    for j in fullText:
        #remove newlines
        j = j.replace('\n\n','\n')
        new.add_paragraph(j)
                
    new.save(doc2)

#get arguments so we can say which method to call from command line
#first arg should be del, order, or font, to call the method
#second and third args should be 
def main():
    #check for python 3 or 2.x
    py3 = sys.version_info[0] > 2

    #message prompt for method
    mtdmsg = """If you would like to reorder the text, type: order;
    to remove blank lines, type: del; 
    to change font size, type: font. \n"""

    if py3:
        mtd = input(mtdmsg)
        doc1 = input("what is the name of the file you would like to change?\n")
        doc2 = input("what should the name of the file be after changing?\n")
        if mtd == 'order':
            reOrder(doc1, doc2)
            print('finished!')
        elif mtd == 'del':
            delEmpty(doc1, doc2)
            print('finished!')
        elif mtd == 'font':
            fontSize(doc1, doc2)
            print('finished!')
        else:
            print('something went wrong, please try again and make sure you choose one of the 3 commands.')
    else:
        mtd = raw_input(mtdmsg)
        doc1 = raw_input("what is the name of the file you would like to change?\n")
        doc2 = raw_input("what should the name of the file be after changing?\n")
        if mtd == 'order':
            reOrder(doc1, doc2)
            print('finished!')
        elif mtd == 'del':
            delEmpty(doc1, doc2)
            print('finished!')
        elif mtd == 'font':
            fontSize(doc1, doc2)
            print('finished!')
        else:
            print('something went wrong, please try again and make sure you choose one of the 3 commands.')




if __name__ == '__main__':
    main()


