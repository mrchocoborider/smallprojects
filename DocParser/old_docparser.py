import re
import docx
from docx.shared import Pt
import copy


reg = re.compile(r'[a-zA-Z]+')

#Most of these methods will have to be customized to new documents
#Be aware of that fact!

def reOrder(docname):
    #docname should be a string
    doc = docx.Document(docname)

    fullText = []
    ctxt = []
    etxt = []

    #make a variable called lastline to see if the last line
    #was English or Chinese (they need to be kept together)
    lastLine = ""

    #we're going to switch the order of English and Chinese lines
    k = 0
    for i in doc.paragraphs:
        #First two lines should be appended without change
        if k == 0 or k == 1:
            fullText.append(i)
        if k != 0 and k != 1:
            #English
            if reg.search(i.text) != None:
                etxt.append(i)     
                thisLine = "eng"
                
            #Chinese
            if reg.search(i.text) == None:
                ctxt.append(i)
                thisLine = "chi"
                #for r in i.runs:
                #    print r.font.name

            if lastLine == "" and thisLine == "chi":
                pass
            elif lastLine == "" and thisLine == "eng":
                fullText.append(etxt[0])
                fullText.append(ctxt[0])
                etxt = []
                ctxt = []
                lastLine = "eng"
            elif lastLine == "eng" and thisLine == "chi":
                fullText.append(ctxt[0])
                ctxt = []
                lastLine = "chi"
            elif lastLine == "chi" and thisLine == "chi":
                fullText.append(ctxt[0])
                ctxt = []
                lastLine = "chi"
            elif lastLine == "chi" and thisLine == "eng":
                fullText.append(etxt[0])
                etxt = []
                lastLine = "eng"
            elif lastLine == "eng" and thisLine == "eng":
                fullText.append(etxt[0])
                etxt = []
                lastLine = "eng"


        k += 1

    #blah = '\n'.join(fullText)
    #print blah

    new = docx.Document()

    for j in fullText:
        new.add_paragraph(j.text)
    #an attempt at copying the paragraph objects directly
    #for j in fullText:
    #    new.paragraphs.append(copy.deepcopy(j))

    #print new.paragraphs

    new.save('testing.docx')


def fontSize(docname):
    #docname should be a string
    doc = docx.Document(docname)
    #this is for changing font size of English only lines
    
    for i in doc.paragraphs:
        #if reg.search(i.text) != None:
            #we have to access all the run objects to get to font size
            #for the first doc, each line was a new paragraph
            #another doc needs to break it down to runs
            for j in i.runs:
                if reg.search(j.text) != None:
                    #print j.text
                    j.font.size = Pt(11)


    doc.save('vanessa2.docx')        

def delEmpty(docname):
    #This is written for a specific file that has some weird 
    #characteristics. All the text is organized into just eleven
    #paragraphs
    doc = docx.Document(docname)
    new = docx.Document()

    fullText = []

    #print len(doc.paragraphs)
    for i in doc.paragraphs:
        fullText.append(i.text)

    for j in fullText:
        #print j
        #remove newlines
        j = j.replace('\n\n','\n')
        #print j
        new.add_paragraph(j)

    new.save('docname.docx')

reOrder('reorder.docx')
#fontSize('testing.docx')
#delEmpty('vanessa.docx')
